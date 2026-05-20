"""
/lineup router — game day position rotation generator
POST /lineup/generate  → computes rotation, returns docx_url
GET  /lineup/download/{filename} → serves the generated file
GET  /lineup/options → valid form choices
"""

import uuid
from pathlib import Path
from typing import List, Optional, Dict
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(__file__).parent.parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

router = APIRouter()


# ── Pydantic models ─────────────────────────────────────────────────────────

class PlayerIn(BaseModel):
    name: str
    number: Optional[str] = None


class LineupRequest(BaseModel):
    date: str
    team_name: str = "Team"
    innings: int = 6
    batting_order: List[PlayerIn]
    pitchers: List[str]
    catchers: List[str]
    outfield_format: str = "standard"  # "standard" | "4-outfielder"


class LineupResponse(BaseModel):
    docx_url: str


# ── Position sets ────────────────────────────────────────────────────────────

FIELD_POSITIONS = {
    "standard":       ["1B", "2B", "3B", "SS", "LF", "CF", "RF"],
    "4-outfielder":   ["1B", "2B", "3B", "SS", "LF", "LC", "RC", "RF"],
    "8p":             ["1B", "2B", "3B", "SS", "LF", "RF"],
}

DISPLAY_POSITIONS = {
    "standard":       ["P", "C", "1B", "2B", "3B", "SS", "LF", "CF", "RF"],
    "4-outfielder":   ["P", "C", "1B", "2B", "3B", "SS", "LF", "LC", "RC", "RF"],
    "8p":             ["P", "C", "1B", "2B", "3B", "SS", "LF", "RF"],
}

POS_ORDERS_STANDARD = [
    ["1B", "2B", "3B", "SS", "LF", "CF", "RF"],
    ["RF", "CF", "LF", "SS", "3B", "2B", "1B"],
    ["SS", "3B", "2B", "1B", "LF", "CF", "RF"],
    ["LF", "CF", "RF", "SS", "1B", "3B", "2B"],
    ["3B", "SS", "CF", "1B", "RF", "LF", "2B"],
    ["2B", "1B", "SS", "3B", "LF", "RF", "CF"],
]

POS_ORDERS_4OF = [
    ["1B", "2B", "3B", "SS", "LF", "LC", "RC", "RF"],
    ["RF", "RC", "LC", "LF", "SS", "3B", "2B", "1B"],
    ["LC", "RC", "LF", "RF", "SS", "1B", "3B", "2B"],
    ["SS", "3B", "2B", "1B", "RF", "LF", "LC", "RC"],
    ["LF", "RF", "LC", "RC", "1B", "SS", "2B", "3B"],
    ["3B", "SS", "RC", "LC", "1B", "RF", "2B", "LF"],
    ["2B", "1B", "SS", "3B", "LC", "RF", "LF", "RC"],
    ["RC", "LF", "SS", "2B", "RF", "1B", "LC", "3B"],
]


# ── Algorithm ────────────────────────────────────────────────────────────────

def try_assign(pos_idx, assigned, pos_left, players, result, history):
    if pos_idx == len(pos_left):
        return result
    pos = pos_left[pos_idx]
    available = [p for p in players if p not in assigned]
    available.sort(key=lambda p: (1 if pos in history[p] else 0, len(history[p])))
    for player in available:
        final = try_assign(
            pos_idx + 1, assigned | {player}, pos_left, players,
            {**result, pos: player}, history,
        )
        if final is not None:
            return final
    return None


def plan_bench(all_players, innings, fixed, bench_per_inning):
    if bench_per_inning <= 0:
        return [[] for _ in range(innings)]

    duties = {p: 0 for p in all_players}
    for f in fixed:
        duties[f["pitcher"]] += 1
        duties[f["catcher"]] += 1

    total_slots = bench_per_inning * innings
    base = total_slots // len(all_players)
    extra = total_slots % len(all_players)
    sorted_by_duty = sorted(all_players, key=lambda p: duties[p])
    targets = {p: base + (1 if i < extra else 0) for i, p in enumerate(sorted_by_duty)}

    bench = [[] for _ in range(innings)]
    counts = {p: 0 for p in all_players}

    for i in range(innings):
        pitcher, catcher = fixed[i]["pitcher"], fixed[i]["catcher"]
        available = [
            p for p in all_players
            if p != pitcher and p != catcher and counts[p] < targets[p]
        ]
        available.sort(key=lambda p: (duties[p], counts[p]))
        for j in range(bench_per_inning):
            if j < len(available):
                bench[i].append(available[j])
                counts[available[j]] += 1
        if len(bench[i]) < bench_per_inning:
            fallback = sorted(
                [p for p in all_players if p != pitcher and p != catcher and p not in bench[i]],
                key=lambda p: counts[p],
            )
            for fb in fallback:
                if len(bench[i]) >= bench_per_inning:
                    break
                bench[i].append(fb)
                counts[fb] += 1

    return bench


def compute_lineup(req: LineupRequest):
    all_players = [p.name for p in req.batting_order]
    use_4of = req.outfield_format == "4-outfielder"
    use_8p = len(all_players) == 8 and not use_4of
    field_key = "4-outfielder" if use_4of else ("8p" if use_8p else "standard")
    positions_per_inning = 10 if use_4of else (8 if use_8p else 9)
    bench_per_inning = max(0, len(all_players) - positions_per_inning)
    pos_orders = POS_ORDERS_4OF if use_4of else POS_ORDERS_STANDARD
    field_pos = set(FIELD_POSITIONS[field_key])

    fixed = [{"pitcher": req.pitchers[i], "catcher": req.catchers[i]} for i in range(req.innings)]
    bench = plan_bench(all_players, req.innings, fixed, bench_per_inning)

    best_assignments = None
    best_repeats = float("inf")
    best_history = None

    for pos_order in pos_orders:
        hist = {p: set() for p in all_players}
        asgn = []
        ok = True

        for i in range(req.innings):
            pitcher, catcher = fixed[i]["pitcher"], fixed[i]["catcher"]
            bench_this = bench[i]
            field_players = [p for p in all_players if p != pitcher and p != catcher and p not in bench_this]
            ordered_pos = [p for p in pos_order if p in field_pos]
            result = try_assign(0, set(), ordered_pos, field_players, {}, hist)
            if result is None:
                ok = False
                break
            hist[pitcher].add("P")
            hist[catcher].add("C")
            for b in bench_this:
                hist[b].add("BENCH")
            for pos, player in result.items():
                hist[player].add(pos)
            asgn.append(result)

        if not ok:
            continue

        full_history = {p: [] for p in all_players}
        for i in range(req.innings):
            full_history[fixed[i]["pitcher"]].append("P")
            full_history[fixed[i]["catcher"]].append("C")
            for b in bench[i]:
                full_history[b].append("BENCH")
            for pos, name in asgn[i].items():
                full_history[name].append(pos)

        repeats = 0
        for positions in full_history.values():
            playing = [p for p in positions if p != "BENCH"]
            for j, p in enumerate(playing):
                if playing.index(p) != j:
                    repeats += 1

        if repeats < best_repeats:
            best_repeats = repeats
            best_assignments = asgn
            best_history = full_history
            if repeats == 0:
                break

    return best_assignments, best_history, bench, bench_per_inning, field_key


# ── docx generation ──────────────────────────────────────────────────────────

NAVY    = "1B4F8C"
ALT     = "D9E8F5"
WHITE   = "FFFFFF"
RED     = "C00000"
GREEN   = "1F6B2E"
GRAY    = "888888"
BENCHBG = "EFEFEF"
DARK    = "222222"


def _rgb(hex_str):
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell(table_cell, text, bold=False, italic=False, center=True,
          fg=DARK, bg=None, size=18):
    table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = table_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size / 2)
    run.font.color.rgb = _rgb(fg)
    run.font.name = "Arial"
    if bg:
        _set_cell_bg(table_cell, bg)


def _pos_style(pos):
    if pos == "P":     return dict(fg=RED,   bold=True,  italic=False, bg=None)
    if pos == "C":     return dict(fg=GREEN, bold=True,  italic=False, bg=None)
    if pos == "BENCH": return dict(fg=GRAY,  bold=False, italic=True,  bg=BENCHBG)
    return dict(fg=DARK, bold=False, italic=False, bg=None)


def _disp(name, roster_map):
    num = roster_map.get(name)
    return f"{name} #{num}" if num else name


def generate_docx(req: LineupRequest, assignments, bench, bench_per_inning, field_key):
    roster_map = {p.name: p.number for p in req.batting_order if p.number}
    all_players = [p.name for p in req.batting_order]
    innings = req.innings
    display_pos = DISPLAY_POSITIONS[field_key]
    if bench_per_inning > 0:
        display_pos = display_pos + ["BENCH"]
    has_bench = bench_per_inning > 0
    fixed = [{"pitcher": req.pitchers[i], "catcher": req.catchers[i]} for i in range(innings)]

    doc = Document()
    section = doc.sections[0]
    section.orientation = 1  # landscape
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin   = Inches(0.6)
    section.right_margin  = Inches(0.6)
    section.top_margin    = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"⚾ {req.team_name} — Game Day Lineup | {req.date}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _rgb("1B3A6B")
    run.font.name = "Arial"

    # Subtitle
    bench_note = f"{bench_per_inning} bench/inning" if has_bench else "no bench"
    sub_fmt = "4-outfielder" if req.outfield_format == "4-outfielder" else "standard outfield"
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"{len(all_players)} players · {innings} innings · {sub_fmt} · {bench_note}"
    )
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = _rgb("555555")
    sub_run.font.name = "Arial"

    doc.add_paragraph()

    # ── Table 1: By Position ─────────────────────────────────────────────────
    heading = doc.add_paragraph("By Position")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(11)
    heading.runs[0].font.color.rgb = _rgb("1B3A6B")

    col_count = 1 + innings
    t1 = doc.add_table(rows=1 + len(display_pos), cols=col_count)
    t1.style = "Table Grid"

    # Header row
    _cell(t1.rows[0].cells[0], "Position", bold=True, fg=WHITE, bg=NAVY)
    for j in range(innings):
        _cell(t1.rows[0].cells[j + 1], f"Inn {j + 1}", bold=True, fg=WHITE, bg=NAVY)

    for row_idx, pos in enumerate(display_pos):
        row = t1.rows[row_idx + 1]
        row_bg = ALT if row_idx % 2 == 0 else WHITE
        _cell(row.cells[0], pos, bold=True, center=False, fg="444444", bg=row_bg)

        for j in range(innings):
            a = assignments[j]
            if pos == "P":
                text = _disp(fixed[j]["pitcher"], roster_map)
                style = _pos_style("P")
            elif pos == "C":
                text = _disp(fixed[j]["catcher"], roster_map)
                style = _pos_style("C")
            elif pos == "BENCH":
                text = ", ".join(_disp(b, roster_map) for b in bench[j])
                style = _pos_style("BENCH")
            else:
                player = a.get(pos, "")
                text = _disp(player, roster_map) if player else "—"
                style = _pos_style(pos)

            cell_bg = style.get("bg") or row_bg
            _cell(row.cells[j + 1], text,
                  bold=style["bold"], italic=style["italic"],
                  fg=style["fg"], bg=cell_bg)

    doc.add_paragraph()

    # ── Table 2: By Player ───────────────────────────────────────────────────
    heading2 = doc.add_paragraph("By Player")
    heading2.runs[0].bold = True
    heading2.runs[0].font.size = Pt(11)
    heading2.runs[0].font.color.rgb = _rgb("1B3A6B")

    t2 = doc.add_table(rows=1 + len(all_players), cols=2 + innings)
    t2.style = "Table Grid"

    _cell(t2.rows[0].cells[0], "#",      bold=True, fg=WHITE, bg=NAVY)
    _cell(t2.rows[0].cells[1], "Player", bold=True, fg=WHITE, bg=NAVY, center=False)
    for j in range(innings):
        _cell(t2.rows[0].cells[j + 2], f"Inn {j + 1}", bold=True, fg=WHITE, bg=NAVY)

    # Build position lookup per player per inning
    player_positions: Dict[str, List[str]] = {p: [] for p in all_players}
    for i in range(innings):
        player_positions[fixed[i]["pitcher"]].append("P")
        player_positions[fixed[i]["catcher"]].append("C")
        for b in bench[i]:
            player_positions[b].append("BENCH")
        for pos, name in assignments[i].items():
            player_positions[name].append(pos)

    for row_idx, player in enumerate(req.batting_order):
        row = t2.rows[row_idx + 1]
        row_bg = ALT if row_idx % 2 == 0 else WHITE
        _cell(row.cells[0], str(row_idx + 1), fg="888888", bg=row_bg)
        _cell(row.cells[1], _disp(player.name, roster_map),
              bold=True, center=False, fg=DARK, bg=row_bg)

        positions = player_positions.get(player.name, [])
        for j, pos in enumerate(positions):
            style = _pos_style(pos)
            cell_bg = style.get("bg") or row_bg
            _cell(row.cells[j + 2], pos,
                  bold=style["bold"], italic=style["italic"],
                  fg=style["fg"], bg=cell_bg)

    # Legend footnote
    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legend_run = legend.add_run(
        "P = Pitcher  ·  C = Catcher  ·  BENCH = Sitting out"
        + (f"  ·  {req.team_name} — {req.date}" if req.team_name else "")
    )
    legend_run.font.size = Pt(8)
    legend_run.font.color.rgb = _rgb("888888")
    legend_run.font.name = "Arial"

    filename = f"lineup_{uuid.uuid4().hex[:8]}.docx"
    path = OUTPUT_DIR / filename
    doc.save(str(path))
    return str(path)


# ── Endpoints ────────────────────────────────────────────────────────────────

@router.post("/generate", response_model=LineupResponse)
def generate_lineup(req: LineupRequest):
    if len(req.batting_order) < 8:
        raise HTTPException(status_code=422, detail="Minimum 8 players required")
    if len(req.pitchers) != req.innings or len(req.catchers) != req.innings:
        raise HTTPException(status_code=422, detail="pitchers and catchers must each have one entry per inning")

    for i in range(req.innings):
        if req.pitchers[i] == req.catchers[i]:
            raise HTTPException(
                status_code=422,
                detail=f"Inning {i + 1}: {req.pitchers[i]} cannot be both pitcher and catcher",
            )

    try:
        assignments, _, bench, bench_per_inning, field_key = compute_lineup(req)
        if assignments is None:
            raise HTTPException(status_code=500, detail="Could not compute a valid rotation — check inputs")

        path = generate_docx(req, assignments, bench, bench_per_inning, field_key)
        filename = Path(path).name
        return LineupResponse(docx_url=f"/api/lineup/download/{filename}")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/download/{filename}")
def download_file(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path=str(path),
        filename=filename,
        media_type="application/octet-stream",
    )


@router.get("/options")
def get_options():
    return {
        "innings":        [4, 5, 6],
        "outfield_count": [3, 4],
    }
