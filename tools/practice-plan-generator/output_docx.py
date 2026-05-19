"""
output_docx.py — Generate a .docx practice plan (v2 — 3-phase structure)
"""
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path(__file__).parent / "outputs"

NAVY     = "1B3A6B"
RED      = "C8102E"
LT_BLUE  = "E8EFF8"
ALT_ROW  = "F5F7FA"
WHITE    = "FFFFFF"
DARK     = "1A1A2E"
GRAY     = "888888"
MED_BLUE = "2E5FA3"
GREEN_BG = "E8F5E9"
AMBER_BG = "FFF8E1"
PURPLE_BG= "F3E5F5"

PHASE_COLORS = {
    "opening":  (LT_BLUE,  NAVY),
    "stations": (GREEN_BG, "1B5E20"),
    "team":     (AMBER_BG, "E65100"),
    "closure":  (LT_BLUE,  NAVY),
}

GAME_LIKE_TEAM_IDS = {
    "two_pitch_game", "game_situation_scrimmage", "live_bp_rotation",
    "live_bp_wiffle_rotation", "fielding_doubles_game", "fielding_singles_game",
    "infield_outfield_combined",
}

def generate_docx(plan: dict) -> str:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.7)
        sec.left_margin = sec.right_margin = Inches(0.8)

    def shade_para(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def h_rgb(hex6):
        r, g, b = int(hex6[0:2],16), int(hex6[2:4],16), int(hex6[4:6],16)
        return RGBColor(r, g, b)

    def add_para(text, size=10, bold=False, italic=False, color=DARK,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
                 shade=None, indent=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        if shade:
            shade_para(p, shade)
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.font.color.rgb = h_rgb(color)
        return p

    # ── Title ─────────────────────────────────────────────────────────────────
    add_para("LGLL PRACTICE PLAN", size=16, bold=True, color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, shade=NAVY,
             space_before=6, space_after=2)
    add_para("La Grange Little League  ·  Ages 7–9", size=10, color="CCDAF0",
             align=WD_ALIGN_PARAGRAPH.CENTER, shade=NAVY, space_after=8)

    # ── Meta table ────────────────────────────────────────────────────────────
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    labels = ["DATE", "LOCATION", "DURATION / COACHES", "FOCUS"]
    values = [plan["date"], plan["location"],
              f"{plan['duration']} min  ·  {plan['coaches']} coaches  ·  ~{plan['player_count']} players",
              plan["focus"]]
    widths = [Inches(1.2), Inches(2.5), Inches(1.8), Inches(2.0)]
    for ci, (lbl, val) in enumerate(zip(labels, values)):
        hc = meta.rows[0].cells[ci]; hc.width = widths[ci]
        shade_cell(hc, NAVY)
        hr = hc.paragraphs[0].add_run(lbl)
        hr.bold=True; hr.font.size=Pt(8); hr.font.name="Arial"
        hr.font.color.rgb=h_rgb(WHITE)
        hc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

        vc = meta.rows[1].cells[ci]
        shade_cell(vc, ALT_ROW)
        vr = vc.paragraphs[0].add_run(val)
        vr.font.size=Pt(9); vr.font.name="Arial"; vr.font.color.rgb=h_rgb(DARK)
        vc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Phase legend strip ────────────────────────────────────────────────────
    legend_p = doc.add_paragraph()
    shade_para(legend_p, "F8F9FA")
    legend_p.paragraph_format.space_before = Pt(2)
    legend_p.paragraph_format.space_after  = Pt(6)
    legend_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, color in [("OPENING", NAVY), ("  STATIONS", "1B5E20"),
                          ("  TEAM TIME", "E65100"), ("  CLOSURE", NAVY)]:
        lr = legend_p.add_run(f"  {label}  ")
        lr.bold=True; lr.font.size=Pt(8); lr.font.name="Arial"
        lr.font.color.rgb=h_rgb(color)

    # ── Phases ────────────────────────────────────────────────────────────────
    running_clock = 0

    for phase in plan["phases"]:
        bg_color, text_color = PHASE_COLORS.get(phase["type"], (LT_BLUE, NAVY))
        t   = phase["start"]
        end = t + phase["minutes"]

        # Phase header
        header_text = (
            f"  {phase['label'].upper()}  ·  {phase['minutes']} min  ·  T+{t:02d}–{end:02d}"
        )
        if phase["type"] == "stations":
            sd = phase["data"]
            header_text += f"  ·  {sd['n_rotations']} stations × {sd['rotation_minutes']} min"

        add_para(header_text, size=11, bold=True, color=text_color,
                 shade=bg_color, space_before=8, space_after=2)

        if phase["type"] == "opening":
            cur_sub = None
            for item in phase["items"]:
                d = item["drill"]
                if item["sub_phase"] != cur_sub:
                    cur_sub = item["sub_phase"]
                    add_para(f"  ▸ {cur_sub}", size=9, bold=True, color=MED_BLUE,
                             shade=WHITE, space_before=4, space_after=1)

                # Drill row
                tbl = doc.add_table(rows=1, cols=2)
                tbl.style = "Table Grid"
                tc = tbl.rows[0].cells[0]; tc.width = Inches(1.1)
                shade_cell(tc, ALT_ROW)
                tr = tc.paragraphs[0].add_run(f"{item['minutes']} min")
                tr.bold=True; tr.font.size=Pt(10); tr.font.name="Arial"
                tr.font.color.rgb=h_rgb(RED)
                tc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

                dc = tbl.rows[0].cells[1]
                shade_cell(dc, WHITE)
                nr = dc.paragraphs[0].add_run(d["name"])
                nr.bold=True; nr.font.size=Pt(10); nr.font.name="Arial"
                nr.font.color.rgb=h_rgb(DARK)
                for cue in d.get("coaching_cues", []):
                    cp = dc.add_paragraph()
                    cr = cp.add_run(f"  ▸ {cue}")
                    cr.font.size=Pt(8); cr.font.italic=True; cr.font.name="Arial"
                    cr.font.color.rgb=h_rgb(MED_BLUE)

        elif phase["type"] == "stations":
            sd = phase["data"]
            add_para(f"  Rotate every {sd['rotation_minutes']} minutes on the whistle.",
                     size=9, italic=True, color=GRAY, shade=WHITE,
                     space_before=2, space_after=4)

            letters = "ABCDE"
            for i, s in enumerate(sd["stations"]):
                d = s["drill"]
                coach_note = "coach present" if s["coach_assigned"] else "self-directed"
                bg = ALT_ROW if i % 2 == 0 else WHITE

                # Station header row
                tbl = doc.add_table(rows=0, cols=3)
                tbl.style = "Table Grid"

                # Header
                hrow = tbl.add_row()
                hrow.cells[0].width = Inches(0.6)
                hrow.cells[1].width = Inches(2.5)
                hrow.cells[2].width = Inches(4.4)
                for ci2, (txt, clr) in enumerate([
                    (f"Stn {letters[i]}", WHITE),
                    (s["label"], WHITE),
                    (f"{coach_note}  ·  ~{s['group_size']} players  ·  {s['minutes']} min", WHITE),
                ]):
                    shade_cell(hrow.cells[ci2], MED_BLUE)
                    r2 = hrow.cells[ci2].paragraphs[0].add_run(txt)
                    r2.bold=True; r2.font.size=Pt(9); r2.font.name="Arial"
                    r2.font.color.rgb=h_rgb(WHITE)

                # Drill name + description row
                drow = tbl.add_row()
                shade_cell(drow.cells[0], bg)
                shade_cell(drow.cells[1], bg)
                shade_cell(drow.cells[2], bg)

                nr = drow.cells[1].paragraphs[0].add_run(d["name"])
                nr.bold=True; nr.font.size=Pt(10); nr.font.name="Arial"
                nr.font.color.rgb=h_rgb(DARK)
                if d.get("level") == "intermediate":
                    sr = drow.cells[1].paragraphs[0].add_run("  ★")
                    sr.font.size=Pt(8); sr.font.color.rgb=h_rgb(RED); sr.font.name="Arial"

                desc = d.get("description","").replace("\n"," ").strip()
                dr2 = drow.cells[2].paragraphs[0].add_run(desc[:200])
                dr2.font.size=Pt(9); dr2.font.name="Arial"; dr2.font.color.rgb=h_rgb("333333")
                for cue in d.get("coaching_cues", []):
                    cp = drow.cells[2].add_paragraph()
                    cr = cp.add_run(f"▸ {cue}")
                    cr.font.size=Pt(8); cr.font.italic=True; cr.font.name="Arial"
                    cr.font.color.rgb=h_rgb(NAVY)

                doc.add_paragraph()

        elif phase["type"] == "team":
            for item in phase["items"]:
                d   = item["drill"]
                tag = "GAME-LIKE" if d["id"] in GAME_LIKE_TEAM_IDS else "INSTRUCTIVE"
                tag_color = "1B5E20" if tag == "GAME-LIKE" else "E65100"

                tbl = doc.add_table(rows=1, cols=2)
                tbl.style = "Table Grid"
                tc = tbl.rows[0].cells[0]; tc.width = Inches(1.2)
                shade_cell(tc, ALT_ROW)
                tr = tc.paragraphs[0].add_run(f"{item['minutes']} min")
                tr.bold=True; tr.font.size=Pt(11); tr.font.name="Arial"
                tr.font.color.rgb=h_rgb(RED)
                tc.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                tag_r = tc.add_paragraph().add_run(tag)
                tag_r.bold=True; tag_r.font.size=Pt(7); tag_r.font.name="Arial"
                tag_r.font.color.rgb=h_rgb(tag_color)
                tc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

                dc = tbl.rows[0].cells[1]
                shade_cell(dc, WHITE)
                nr = dc.paragraphs[0].add_run(d["name"])
                nr.bold=True; nr.font.size=Pt(11); nr.font.name="Arial"
                nr.font.color.rgb=h_rgb(DARK)
                desc = d.get("description","").replace("\n"," ").strip()
                dp = dc.add_paragraph()
                dr2 = dp.add_run(desc[:250])
                dr2.font.size=Pt(9); dr2.font.name="Arial"; dr2.font.color.rgb=h_rgb("333333")
                for cue in d.get("coaching_cues", []):
                    cp = dc.add_paragraph()
                    cr = cp.add_run(f"▸ {cue}")
                    cr.font.size=Pt(9); cr.font.italic=True; cr.font.name="Arial"
                    cr.font.color.rgb=h_rgb(NAVY)

                doc.add_paragraph()

        elif phase["type"] == "closure":
            add_para(
                "  Bring team in. Highlight 2-3 specific positives from today. "
                "Name one focus area for next practice. End with a team cheer.",
                size=9, color=DARK, shade=WHITE, indent=0.2,
                space_before=2, space_after=6
            )

    # ── Equipment ─────────────────────────────────────────────────────────────
    add_para("  EQUIPMENT NEEDED", size=10, bold=True, color=NAVY,
             shade=LT_BLUE, space_before=8, space_after=2)
    eq_items = [e.replace("_"," ").title() for e in plan["all_equipment"]]
    add_para("  " + "  ·  ".join(eq_items) if eq_items else "  No special equipment",
             size=9, color=DARK, shade=WHITE, space_after=4)

    # ── Footer ────────────────────────────────────────────────────────────────
    add_para(
        f"Generated {datetime.now().strftime('%B %d, %Y')}  ·  "
        f"LGLL Practice Plan Generator  ·  {plan['total_minutes']} / {plan['duration']} min",
        size=7, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8
    )

    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"practice_plan_{date_str}_{plan['location_key']}_{plan['duration']}min.docx"
    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))
    return str(out_path)
